import re

import regex
from retry import retry
from pathlib import Path

PDF_PATTERN = r".*\.pdf"
DOCX_PATTERN = r".*\.docx"
DECIMAL_PATTERN = r"\d+\.\s"


class QuestionsParser:
    def __init__(self, path_to_pdf, remainder_mod: tuple = (0, 1), questions_range: tuple = None):
        """
        The __init__ function is called when an instance of the class is created.

        :param path_to_pdf: Store the path to the pdf file
        :param remainder_mod: Tuple containing the remainder and the divisor respectively
        """
        import cloudscraper

        print(_validate_path(path_to_pdf))
        self.__question_list = _parse_questions(_validate_path(path_to_pdf))

        self.__remainder_mod = remainder_mod
        self.__questions_range = questions_range
        self.__scraper = cloudscraper.create_scraper()
        self.headers = {
                "User-Agent":
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:104.0) Gecko/20100101 Firefox/104.0"
            }
        self.params = {
            "hl": "en",
        }
        self.__result = None

    @retry(exceptions=ConnectionError, tries=3, delay=1, backoff=0)
    def parse_google(self, autoparse_answers: bool = True):
        """
        Takes in a list of questions and returns a dictionary with the question as the key
        and another dictionary as its value. The nested dictionary contains two keys, 'link' and 'answer'. The link is
        the url of where to find an answer to that specific question on Google.com, while the answer is what Google
        tells us is an answer to that question.

        :param autoparse_answers: If True, the answers will be parsed automatically. Default is True
        :param self: Reference the object itself in a method
        :return: QuestionsParser object
        """
        import bs4
        from QuestionsParser.utils.progress_bar import progress_bar

        def parse_link(_soup):
            """
            Parse the link from the soup
            """
            try:
                _link = _soup.find("div", {"class": "yuRUbf"}).find("a")["href"]
            except AttributeError:
                raise ConnectionError("Something went wrong. Please try again later")
            return _link

        def parse_answer(_soup):
            possible_answer = soup.find("span", {"class": "hgKElc"})
            if possible_answer:
                return {"text": possible_answer.text, "confidence": 1}

            possible_answer = soup.find("div", {"class": "VwiC3b yXK7lf MUxGbd yDYNvb lyLwlc lEBKkf"})
            if possible_answer:
                result[question_str]["answer"] = {"text": possible_answer.text, "confidence": 0.5}
                text = possible_answer.text.rsplit('. ', 1)[0]
                return {"text": text, "confidence": 0.5}
            else:
                return {"text": "No answer found", "confidence": 0}

        result = dict()

        for i, question in enumerate(progress_bar(self.__question_list)):
            if (i + 1) % self.__remainder_mod[1] != self.__remainder_mod[0]:
                continue
            if self.__questions_range and i + 1 < self.__questions_range[0]:
                continue
            if self.__questions_range and i + 1 > self.__questions_range[1]:
                continue

            question_str = str(i + 1) + ". " + question
            result[question_str] = dict()

            question = question.replace("?", "") + " in Java?"

            request_result = self.__scraper.get(
                f"https://www.google.com/search?q={question}",
                headers=self.headers,
                params=self.params
            )

            soup = bs4.BeautifulSoup(request_result.text, "html.parser")

            link = parse_link(soup)
            result[question_str]["link"] = link

            if autoparse_answers:
                answer = parse_answer(soup)
                result[question_str]["answer"] = answer

        self.__result = result
        return self

    def write_to_file(self, path_to_file, header: str = "Answers"):
        """
        Write the result to a file. The file will be overwritten if it already exists.
        Note: The only file extensions supported are .pdf and .docx. If possible, use .docx, as it is more readable
        and convenient.
        """
        if self.__result is None:
            raise ValueError("No result found. Please call parse_answers_google() first")
        if regex.match(PDF_PATTERN, path_to_file):
            _write_to_pdf(path_to_file, self.__result, header)
        elif regex.match(DOCX_PATTERN, path_to_file):
            _write_to_docx(path_to_file, self.__result, header)
        else:
            raise ValueError("Only .pdf and .docx files are supported")

    def get_questions(self):
        """
        Return the list of questions
        """
        return self.__question_list


def _validate_path(path_to_pdf) -> Path:
    """
    Checks that the path to a pdf file is valid.
    If any of these tests fail, _validate_path raises an error.

    :param path_to_pdf: Check if the path is a valid pdf file
    :return: A boolean value
    """
    if not isinstance(path_to_pdf, str):
        raise TypeError("Path to pdf must be a string")
    if not regex.match(PDF_PATTERN, path_to_pdf):
        raise ValueError("Path to pdf must be a valid path to a pdf file")

    return Path(path_to_pdf)


def _parse_questions(path_to_pdf) -> list[str]:
    """
    The parse_questions function takes a path to a pdf file and returns
    a list of strings, where each string is the text from one question.

    :param path_to_pdf: Specify the path to the pdf file containing all the questions
    :return: A list of strings
    """
    import PyPDF2

    questions_file = open(path_to_pdf, "rb")
    try:
        questions = PyPDF2.PdfFileReader(questions_file).getPage(0).extractText()
    except Exception as e:
        print(e)
        raise ValueError("Invalid PDF file")

    questions_list = [question.strip() for question in re.split(DECIMAL_PATTERN, questions) if question != ""]

    return questions_list


def _write_to_pdf(path_to_file, result, header: str):
    """
    Write the result to a pdf file. The file will be overwritten if it already exists.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=header, ln=1, align="C")

    for question, answer in result.items():
        question = question.encode("latin-1", "replace").decode("latin-1")
        link = answer.get("link")
        answer = answer.get("answer")

        pdf.set_font("Arial", "B", size=13)
        pdf.cell(200, 10, txt=question, ln=1, align="L")
        pdf.set_font("Arial", size=12)

        if link:
            pdf.set_font("Arial", "I", size=12)
            link = link.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(200, 10, txt=link, align="L")
            pdf.set_font("Arial", size=12)
        if answer:
            answer = answer.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(200, 10, txt=answer, align="L")

    pdf.output(path_to_file)
    print(f"Successfully wrote to {path_to_file}")


def _write_to_docx(path_to_file, result, header: str):
    """
    Write the result to a docx file. The file will be overwritten if it already exists.
    """
    from docx import Document
    from docx.shared import RGBColor
    from QuestionsParser.utils.hyperlink import add_hyperlink

    document = Document()
    document.add_heading(header, 0)

    for question, answer in result.items():
        link = answer.get("link")
        answer = answer.get("answer")

        document.add_heading(question, 1)

        if link:
            link_paragraph = document.add_paragraph()
            link_paragraph.add_run("Link: ").bold = True
            add_hyperlink(link_paragraph, link, link)

        if answer:
            if answer["text"] == "No answer found":
                run = document.add_paragraph().add_run("No answer found. Please check the link above.")
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
                continue

            if answer["confidence"] < 1:
                answer = document.add_paragraph().add_run(answer["text"])
                answer.add_break()

                run = document.add_paragraph().add_run("This answer is not very confident. "
                                                       "Please check the link above.")
                run.font.color.rgb = RGBColor(255, 150, 0)
            else:
                document.add_paragraph(answer["text"])

    document.save(path_to_file)
    print(f"Successfully wrote to {path_to_file}")
